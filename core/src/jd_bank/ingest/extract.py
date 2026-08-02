"""Text extraction for the real SFU archive formats.

Ported and slimmed from hris ``pipeline/parsing/extract.py`` (reuse map #14).
What is kept: the tolerant decode ladder, the DOCX (python-docx) and RTF
(striprtf) backends, and the NUL-stripping sanitiser that keeps Postgres happy.
What is dropped: the PDF/PyMuPDF path and the page/block structure — the JD
corpus is near-zero PDF (census §3) and downstream stages want a flat text
stream, not positioned blocks. What is *added*: an ``antiword`` backend for the
legacy binary ``.doc`` corpus (4,577 files; python-docx cannot read binary
``.doc``), which the census validated as the offline extraction path.

The public entry point is :func:`extract_text` — ``bytes + DocumentFormat`` in,
plain ``str`` out. Unsupported formats raise :class:`UnsupportedFormatError`.
"""

from __future__ import annotations

import hashlib
import re
import subprocess
import tempfile
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path
from typing import cast

from docx import Document
from docx import types as docx_types
from docx.blkcntnr import BlockItemContainer
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.jd_bank.db.models import DocumentFormat
from src.jd_core.parser.headings import (
    CANONICAL_HEADING,
    DEPARTMENT_LABEL_RX,
    EMPLOYEE_GROUP_LABEL_RX,
    GRADE_LABEL_RX,
    POSITION_NO_LABEL_RX,
    TITLE_LABEL_RX,
    SectionKey,
)

# antiword ships in the ingestion image (Dockerfile). Overridable for tests.
ANTIWORD_BIN = "antiword"

# DoS hardening. The archive is ~14,565 files processed by shared workers, so a
# single corrupt/crafted document must not hang a worker or exhaust its memory.
# A crafted binary .doc can send antiword into a pathological loop -> hard wall
# clock cap; and no legitimate SFU JD approaches 50 MiB (the largest golden file
# is ~5 MiB) -> reject oversized inputs before loading them into a backend.
#
# NB what the cap protects is THE EXTRACTOR — antiword / python-docx, which allocate
# proportional to their input and are the actual hazard. It does NOT forbid us from
# *hashing* bytes we never parse: see `stream_sha256`, which reads any file in constant
# memory. An oversized file is still a file, and it still gets a `source_documents` row.
ANTIWORD_TIMEOUT_SECONDS = 30
MAX_DOCUMENT_BYTES = 50 * 1024 * 1024  # 50 MiB

#: Chunk size for :func:`stream_sha256`. Bounds the memory a hash of ANY file costs.
_HASH_CHUNK_BYTES = 1024 * 1024  # 1 MiB


class ExtractionError(RuntimeError):
    """Text could not be extracted from an otherwise-supported document."""


class UnsupportedFormatError(ExtractionError):
    """The document format has no extraction backend (route to manual triage)."""


class DocumentTooLargeError(ExtractionError):
    """Input exceeds :data:`MAX_DOCUMENT_BYTES` — rejected before extraction."""


def _strip_nul(text: str) -> str:
    """Remove NUL (U+0000). Postgres ``text``/``jsonb`` cannot store U+0000 (it
    raises ``UntranslatableCharacterError`` at write time), and some legacy
    binaries carry embedded NULs. Cheap unconditional pass — the common clean
    case is a no-op."""
    return text.replace("\x00", "") if "\x00" in text else text


#: UTF-16 byte-order marks (LE, BE). Checked explicitly rather than added to the
#: codec ladder: latin-1 decodes ANY byte sequence, so it can never fall through to a
#: later rung — a UTF-16 file must be recognised up front or it silently becomes
#: mojibake.
_UTF16_BOMS = (b"\xff\xfe", b"\xfe\xff")


def _decode(blob: bytes) -> str:
    """Tolerant byte decode: UTF-16 BOM -> UTF-8 BOM -> UTF-8 -> latin-1 (never
    raises). The UTF-8/latin-1 ladder is hris's, so text inputs decode predictably
    across eras; the UTF-16 rung is ours.

    **Measured defect it fixes:** the archive holds UTF-16LE ``.txt`` exports whose
    ``\\xff\\xfe`` BOM is invalid UTF-8, so the ladder fell through to latin-1 and
    every one of them decoded to ``ÿþP%P%P%…`` — the whole document, title included,
    turned to noise. Small (the archive is 24 ``.txt`` files) but total per file.
    """
    if blob[:2] in _UTF16_BOMS:
        try:
            return blob.decode("utf-16")
        except UnicodeDecodeError:  # truncated/odd-length — fall through
            pass
    for codec in ("utf-8-sig", "utf-8"):
        try:
            return blob.decode(codec)
        except UnicodeDecodeError:
            continue
    return blob.decode("latin-1")


_P = qn("w:p")
_TBL = qn("w:tbl")
_TR = qn("w:tr")
_TC = qn("w:tc")
_SDT = qn("w:sdt")
_SDT_CONTENT = qn("w:sdtContent")


def _iter_docx_block_text(el: BaseOxmlElement) -> Iterator[str]:
    """Walk ``el``'s children in document order, yielding paragraph text.

    ``doc.paragraphs`` (python-docx's high-level API) returns only body-level
    ``<w:p>`` elements — it skips text inside tables (``<w:tbl>``) and Word
    content controls (``<w:sdt>``/``<w:sdtContent>``). This walk descends into
    both, recursively, so arbitrarily nested tables/SDTs (table-in-cell,
    SDT-in-cell, table-in-SDT) are all covered. Elements this loop does not
    recognise (e.g. the trailing ``<w:sectPr>``) are silently skipped — they
    carry no paragraph text.
    """
    for child in el.iterchildren():
        if child.tag == _P:
            # `parent=None` is safe here: `Paragraph.text` only reads
            # `self._p.text` and never touches `parent` (verified against the
            # installed python-docx). `ProvidesStoryPart` is a typing.Protocol
            # python-docx uses for style/part lookups this call path never
            # takes, so the cast is a documented "runtime-safe, type-only" gap
            # rather than a blind ignore.
            yield Paragraph(child, cast(docx_types.ProvidesStoryPart, None)).text
        elif child.tag == _TBL:
            for row in child.findall(_TR):
                for cell in row.findall(_TC):
                    yield from _iter_docx_block_text(cell)
        elif child.tag == _SDT:
            content = child.find(_SDT_CONTENT)
            if content is not None:
                yield from _iter_docx_block_text(content)


#: The heading the recovered header identification block is emitted under. Taken from
#: the parser's own heading vocabulary rather than spelled again here — the flat text
#: this function returns IS the segmenter's input, and its section markers are that
#: interface's structure channel, so a second copy of the word would be a second source
#: of truth for it (the drift `CANONICAL_HEADING` exists to prevent).
IDENTIFICATION_MARKER = CANONICAL_HEADING[SectionKey.IDENTIFICATION]

#: The SAME regexes the segmenter reads identification fields with, so "is this an
#: identification block?" and "can the parser use it?" can never answer differently.
_ID_LABEL_RXS = (
    TITLE_LABEL_RX,
    POSITION_NO_LABEL_RX,
    DEPARTMENT_LABEL_RX,
    EMPLOYEE_GROUP_LABEL_RX,
    GRADE_LABEL_RX,
)


def _is_identification_block(block: str) -> bool:
    """True when ``block`` is an identification table rather than page furniture.

    A title label, or any two distinct identification labels. One lone label is the
    *running* header — ``Position #: <numbers> <page no>`` — which admitting cost 15
    documents a position number of ``2`` (the page number). A real block always names
    the position or states several fields, even when some of them are blank.
    """
    if TITLE_LABEL_RX.search(block):
        return True
    return sum(1 for rx in _ID_LABEL_RXS if rx.search(block)) >= 2


def _collapse(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _row_lines(cells: list[str]) -> Iterator[str]:
    """Pair each ``Label:`` cell with the cell that follows it, one pair per line.

    The identification table is a grid of label/value cells, and the archive packs
    **two pairs onto one row** (measured — 874 of the 876 headers that state a
    grade look like ``Employee Group: | APSA | Grade: | 13``). Joining the whole
    row would leave ``Grade:`` mid-line, where the segmenter's line-anchored
    ``^grade[ \\t]*:`` regex cannot see it; splitting per cell would separate every
    label from its value. Pairing does both correctly.

    A label with no value (an unfilled field) and a value with no label (a banner
    cell) are each emitted alone rather than merged with a neighbour they do not
    belong to.
    """
    index = 0
    while index < len(cells):
        cell = cells[index]
        following = cells[index + 1] if index + 1 < len(cells) else None
        if cell.endswith(":") and following is not None and not following.endswith(":"):
            yield f"{cell} {following}"
            index += 2
        else:
            yield cell
            index += 1


def _iter_container_lines(container: BlockItemContainer) -> Iterator[str]:
    """Yield one line per paragraph, and per table row one line per label/value pair.

    Deliberately different from :func:`_iter_docx_block_text`, which yields each
    table cell on its own line: the segmenter's identification regexes are
    line-anchored (``^Position Title:[ \\t]*(.+)$``), so a label and its value must
    share a line. The body walk keeps cell-per-line because the CUPE/WJQ
    questionnaire parser reads cells as lines.

    Repeated cells from a horizontal merge are collapsed (``row.cells`` yields the
    same cell once per grid column it spans).
    """
    for item in container.iter_inner_content():
        if isinstance(item, Table):
            for row in item.rows:
                cells: list[str] = []
                for cell in row.cells:
                    value = _collapse(cell.text)
                    if value and (not cells or cells[-1] != value):
                        cells.append(value)
                yield from _row_lines(cells)
        else:
            yield _collapse(item.text)


def _docx_identification_block(doc: DocxDocument) -> str:
    """The identification block from a section's header part, or ``""``.

    **Why headers are read at all** (this is the one exception to body-only
    extraction — see ``test_extract_docx_excludes_header_prose_and_all_footers``).
    In the modern SFU template the entire identification table — ``Position
    Title:``, ``Position #:``, ``Department:``, ``Employee Group:``, ``Grade:`` —
    is in ``header*.xml``, not in the body. Measured over all 14,565 archive
    files: **4,968 of 9,948 ``.docx`` carry ``Position Title:`` in the header and
    in no body line**, which is precisely the set whose title parsed as a
    paragraph (the segmenter found no label and fell back to the first content
    line — the About-SFU banner or the Position Summary prose).

    Only a header that :func:`_is_identification_block` recognises is returned; page
    furniture (a running title, a page number) is left excluded, so the exception
    stays as narrow as the evidence. The first-page header is preferred because
    that is where the template puts the full table; the ordinary header usually
    holds just a running ``Position #: … <page>`` line.

    Footers stay excluded entirely: the corpus keeps only revision dates there,
    and there is no heading to file them under, so they would land in whichever
    body section happened to be last (7,171 ``.docx`` have footer text).
    """
    for section in doc.sections:
        for part in (section.first_page_header, section.header):
            lines = [line for line in _iter_container_lines(part) if line]
            if not lines:
                continue
            block = "\n".join(lines)
            if _is_identification_block(block):
                return block
    return ""


def _extract_docx(blob: bytes) -> str:
    """python-docx — document-order body walk. Handles ``.docx`` and
    macro-enabled ``.docm`` (both OOXML).

    Recovers text from tables and Word content controls, which
    ``doc.paragraphs`` alone misses (measured: 2,596 of 9,947 archive
    ``.docx`` lose >40% of their text, 24 lose everything).

    Walks ``doc.element.body``, plus — and only — a header part carrying the
    template's identification labels, which is emitted first under
    :data:`IDENTIFICATION_MARKER` (see :func:`_docx_identification_block`). Header
    prose and all footers remain excluded. A document with no such header extracts
    byte-identically to the body-only walk.
    """
    try:
        doc = Document(BytesIO(blob))
    except Exception as exc:  # noqa: BLE001 — normalise any python-docx failure
        raise ExtractionError(f"docx parse failed: {exc}") from exc
    body = "\n".join(t for t in _iter_docx_block_text(doc.element.body) if t)
    identification = _docx_identification_block(doc)
    if not identification:
        return body
    block = f"{IDENTIFICATION_MARKER}\n{identification}"
    return f"{block}\n{body}" if body else block


def _extract_rtf(blob: bytes) -> str:
    """striprtf — pure-Python RTF -> plain text. Imported lazily so a process
    that never touches RTF need not hard-require the optional dependency."""
    from striprtf.striprtf import rtf_to_text  # noqa: PLC0415 — optional dep, lazy

    try:
        text: str = rtf_to_text(_decode(blob))  # type: ignore[no-untyped-call]
        return text.strip()
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"rtf parse failed: {exc}") from exc


def _extract_txt(blob: bytes) -> str:
    """Plain text — tolerant decode, no structure to recover."""
    return _decode(blob).strip()


def _extract_doc(blob: bytes) -> str:
    """Legacy binary ``.doc`` (Word 97-2003) via the ``antiword`` subprocess.

    python-docx cannot read binary ``.doc``; the census validated antiword as
    the offline path (0 failures across the sample). antiword needs a file on
    disk, so the bytes are written to a temp file, extracted, and cleaned up.
    """
    tmp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(blob)
            tmp_path = tmp.name
        try:
            proc = subprocess.run(
                [ANTIWORD_BIN, tmp_path],
                capture_output=True,
                check=True,
                timeout=ANTIWORD_TIMEOUT_SECONDS,
            )
        except FileNotFoundError as exc:  # antiword not installed
            raise ExtractionError(
                "antiword binary not found — required for legacy .doc extraction"
            ) from exc
        except subprocess.TimeoutExpired as exc:
            raise ExtractionError(
                f"antiword timed out after {ANTIWORD_TIMEOUT_SECONDS}s "
                "(possible corrupt or crafted .doc)"
            ) from exc
        except subprocess.CalledProcessError as exc:
            detail = _decode(exc.stderr or b"").strip()
            raise ExtractionError(f"antiword failed: {detail}") from exc
        return _decode(proc.stdout).strip()
    finally:
        if tmp_path is not None:
            Path(tmp_path).unlink(missing_ok=True)


_BACKENDS = {
    DocumentFormat.DOCX: _extract_docx,
    DocumentFormat.DOC: _extract_doc,
    DocumentFormat.RTF: _extract_rtf,
    DocumentFormat.TXT: _extract_txt,
}


def extract_text(data: bytes, fmt: DocumentFormat) -> str:
    """Extract plain text from ``data`` for the detected ``fmt``.

    ``.docx``/``.docm`` via python-docx, legacy ``.doc`` via antiword, ``.rtf``
    via striprtf, ``.txt`` via a tolerant decode ladder. NUL characters are
    stripped so the result is safe to persist in Postgres. Formats without a
    backend (:data:`DocumentFormat.OTHER` — e.g. ``.pdf``/``.tif``/``.serv``)
    raise :class:`UnsupportedFormatError` for routing to manual triage.
    """
    if len(data) > MAX_DOCUMENT_BYTES:
        raise DocumentTooLargeError(
            f"input is {len(data)} bytes, exceeds cap of {MAX_DOCUMENT_BYTES} bytes"
        )
    backend = _BACKENDS.get(fmt)
    if backend is None:
        raise UnsupportedFormatError(f"no extraction backend for format: {fmt.value}")
    return _strip_nul(backend(data))


def read_document_bytes(path: str | Path) -> bytes:
    """Read a document's bytes, **stat-ing first** so an oversized file is never
    loaded into memory.

    THE ONE HOME OF THE STAT-BEFORE-READ GUARD. :func:`extract_text` also enforces
    :data:`MAX_DOCUMENT_BYTES` — but only *after* the bytes are already in memory, so
    a caller that reads first hands a 50 MiB+ document straight past the DoS guard the
    extractor deliberately added. Every archive walker (the Phase 2.5 baseline runner,
    the Phase 3.2a ingest driver) must go through here rather than calling
    ``path.read_bytes()`` itself.

    **Not hypothetical. MEASURED:** exactly one archive file exceeds the cap —
    ``19980120_19980120_00000293_Asst_to_Director,_Rec_Services.rtf``, **89,397,431
    bytes** — and it sorts inside the FIRST 200 files of the walk, so even a
    ``--limit 200`` smoke test hits it.

    Raises:
        DocumentTooLargeError: the file exceeds :data:`MAX_DOCUMENT_BYTES`. The bytes
            are never read.
        OSError: the file could not be stat-ed or read.
    """
    p = Path(path)
    size = p.stat().st_size
    if size > MAX_DOCUMENT_BYTES:
        raise DocumentTooLargeError(
            f"{size} bytes exceeds the extractor's cap of {MAX_DOCUMENT_BYTES} bytes "
            f"(not read)"
        )
    return p.read_bytes()


def stream_sha256(path: str | Path) -> tuple[str, int]:
    """``(sha256 hex, byte size)`` for a file of ANY size, in **constant memory**.

    The counterpart to :func:`read_document_bytes`, and the thing that lets an
    oversized file still be a *ledger row*. SHA-256 is a streaming digest: the file is
    read in :data:`_HASH_CHUNK_BYTES` chunks and the digest updated per chunk, so the
    89 MB archive ``.rtf`` costs 1 MiB of memory to hash, not 89.

    **Why this exists** (Phase 3.2a, reviewer ruling). The first cut of the ingest
    driver gave an oversized file *no* ``source_documents`` row at all, reasoning: the
    bytes are never read -> there is no sha256 -> the ``(storage_ref, sha256)``
    idempotency key is unsatisfiable. The premise was false. Phase 3.1 made
    ``source_documents`` **one row per FILE — a real ledger** (it took a migration and
    a dropped UNIQUE to get there), and CLAUDE.md non-negotiable #6 is provenance: a
    file that exists in the archive and has no row breaks that property. Today that is
    one file; the property is what matters.

    Hashing is NOT what :data:`MAX_DOCUMENT_BYTES` guards against — the extractor is
    (see the constant). This function never hands a byte to a backend.

    The digest is identical to :func:`~src.jd_bank.ingest.ingest.compute_sha256` over
    the same content (same algorithm, same bytes); ``test_ingest_extract.py`` pins that
    equivalence, because the two must agree or the Tier-1 dedup key would split.
    """
    digest = hashlib.sha256()
    size = 0
    with Path(path).open("rb") as handle:
        while chunk := handle.read(_HASH_CHUNK_BYTES):
            digest.update(chunk)
            size += len(chunk)
    return digest.hexdigest(), size


def extract_text_from_path(path: str | Path) -> str:
    """Convenience wrapper: read a file and extract using its detected format.

    Format detection is by extension (see :func:`src.jd_bank.ingest.ingest.
    detect_format`), imported lazily to avoid a package import cycle.
    """
    from src.jd_bank.ingest.ingest import detect_format  # noqa: PLC0415

    p = Path(path)
    try:
        data = read_document_bytes(p)
    except DocumentTooLargeError as exc:
        # Re-raise NAMING THE FILE. This wrapper's message is human-facing, and it
        # carried the filename before the guard was consolidated.
        #
        # The name is added HERE rather than inside `read_document_bytes` on purpose:
        # that function's message is what the Phase 2.5 baseline writes into its skip
        # ledger, and `docs/baseline/errors.jsonl` is a COMMITTED artifact that must
        # stay byte-identical across runs. Changing the shared message would silently
        # churn it. (The path is already in the ledger's `path` field anyway — it is
        # this wrapper, which callers use without a surrounding row, that needs it in
        # the message.)
        raise DocumentTooLargeError(f"{p.name}: {exc}") from exc
    return extract_text(data, detect_format(p.name))
