"""Render an SFU JD to the official ``.docx`` template (Phase 5.7).

The composer/export half of the plan (§2.5). ``render_sfu_docx`` composes an
:class:`~src.jd_core.models.parsed_jd.SFUJobDescription` into a Word document that
follows SFU's formatting standard (rulebook Part 10.1):

* **Times New Roman, size 10** — set on the ``Normal`` style, so every paragraph
  inherits it.
* **Bold section headers**, in template order; **empty sections are dropped**
  (Part 10.1: "Additional Contextual Info Section: DELETE if not used" — applied to
  every optional section).
* **Standard bullets** for duties, decision-making, problem-solving, relationships
  and qualifications.
* Percentage allocations already live in the duty text as ``(NN%)`` (the composer
  renders them there), so they carry through verbatim.
* The **mandated footer** — the territorial acknowledgement + the Employment-Equity
  statement — is injected from the rulebook boilerplate
  (:func:`~src.jd_core.rules.get_rules`), the single source of that text, never a
  literal here (CLAUDE.md: "footer wording lives in a single config constant").

⚠ The footer wording is a Phase-6 sign-off item — verify it against SFU's current
official text before any external distribution. It is data in ``boilerplate.yaml``,
so that verification is a one-file review, not a code change.

Deterministic and pure of I/O beyond building bytes: same JD (+ rulebook) → same
document. Nothing here validates or publishes — export is a rendering of whatever
draft it is given.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt

from src.jd_core.models.parsed_jd import (
    SFUDuty,
    SFUJobDescription,
    SFUQualification,
    SFURelationships,
)
from src.jd_core.rules import Rules, get_rules

_BODY_FONT = "Times New Roman"
_BODY_SIZE = Pt(10)
_BULLET_STYLE = "List Bullet"


def _set_base_font(doc: DocxDocument) -> None:
    """Times New Roman 10 on the ``Normal`` style — every paragraph inherits it."""
    font = doc.styles["Normal"].font
    font.name = _BODY_FONT
    font.size = _BODY_SIZE


def _header(doc: DocxDocument, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


def _bullets(doc: DocxDocument, header: str, lines: list[str]) -> None:
    """A bold header + one bullet per non-empty line. A no-op if nothing is left after
    dropping blanks, so an empty section never emits a lone header."""
    kept = [line for line in lines if line.strip()]
    if not kept:
        return
    _header(doc, header)
    for line in kept:
        doc.add_paragraph(line, style=_BULLET_STYLE)


def _paragraph_section(doc: DocxDocument, header: str, text: str | None) -> None:
    if not (text or "").strip():
        return
    _header(doc, header)
    doc.add_paragraph(text or "")


def _duty_line(duty: SFUDuty) -> str:
    """The duty as one bullet. The statement usually already leads with the action
    verb (the composer/merge write it that way); prepend the verb only when it does
    not, so a duty never doubles or drops its verb."""
    statement = duty.statement.strip()
    verb = duty.action_verb.strip()
    if verb and not statement.lower().startswith(verb.lower()):
        return f"{verb} {statement}"
    return statement


def _qual_line(qual: SFUQualification) -> str:
    """The qualification as one bullet, with its Toolkit modifier prefixed when it is
    not already part of the text."""
    text = qual.text.strip()
    modifier = (qual.modifier or "").strip()
    if modifier and modifier.lower() not in text.lower():
        return f"{modifier} {text}"
    return text


def _relationships(doc: DocxDocument, rel: SFURelationships | None) -> None:
    if rel is None:
        return
    lines: list[str] = []
    if (rel.supervisory or "").strip():
        lines.append(f"Supervisory: {rel.supervisory}")
    lines += [f"Internal: {c}" for c in rel.internal]
    lines += [f"External: {c}" for c in rel.external]
    _bullets(doc, "Relationships", lines)


def _footer(doc: DocxDocument, rules: Rules) -> None:
    """Inject the mandated territorial acknowledgement + Employment-Equity statement
    from the rulebook boilerplate — always, regardless of the draft's presence flags
    (the footer is required on every published JD)."""
    boilerplate = rules.boilerplate
    _header(doc, "Territorial Acknowledgement & Employment Equity")
    for passage in boilerplate.territorial_acknowledgement:
        doc.add_paragraph(passage)
    for passage in boilerplate.employment_equity:
        doc.add_paragraph(passage)


def render_sfu_docx(jd: SFUJobDescription, *, rules: Rules | None = None) -> bytes:
    """Compose ``jd`` into an SFU-format ``.docx`` and return its bytes.

    Sections render in SFU template order; empty ones are dropped; the mandated
    footer is always appended from the rulebook. Deterministic."""
    rulebook = rules if rules is not None else get_rules()
    doc = Document()
    _set_base_font(doc)

    title = doc.add_paragraph()
    title_run = title.add_run(jd.title)
    title_run.bold = True
    title_run.font.size = Pt(14)

    _paragraph_section(doc, "Position Summary", jd.position_summary)
    _bullets(doc, "Duties and Responsibilities", [_duty_line(d) for d in jd.duties])
    _bullets(doc, "Impact of Decision Making", list(jd.decision_making))
    _bullets(doc, "Problem Solving", list(jd.problem_solving))
    _relationships(doc, jd.relationships)
    _bullets(doc, "Qualifications", [_qual_line(q) for q in jd.qualifications])
    _paragraph_section(doc, "Additional Contextual Information", jd.additional_context)
    _footer(doc, rulebook)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
