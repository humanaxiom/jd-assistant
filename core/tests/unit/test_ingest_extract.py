"""Unit tests for text extraction across the real SFU archive formats.

Covers each backend (txt / rtf / docx / legacy .doc via antiword), the
NUL-strip guard, and the unsupported-format error. The .docx fixture is built
in-process with python-docx (available in the ingestion image) so no binary
blob need be committed; the .doc fixture is one small real archive file
committed under ``tests/fixtures/`` (``.gitattributes`` marks ``*.doc`` binary).
"""

from __future__ import annotations

import subprocess
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

from src.jd_bank.db.models import DocumentFormat
from src.jd_bank.ingest import extract as extract_mod
from src.jd_bank.ingest.extract import (
    IDENTIFICATION_MARKER,
    MAX_DOCUMENT_BYTES,
    DocumentTooLargeError,
    ExtractionError,
    UnsupportedFormatError,
    _iter_docx_block_text,
    extract_text,
    extract_text_from_path,
    stream_sha256,
)
from src.jd_bank.ingest.ingest import compute_sha256
from src.jd_core.parser import parse_jd

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"


def test_stream_sha256_agrees_with_compute_sha256(tmp_path: Path) -> None:
    """The two hashers MUST agree, or Tier-1 dedup splits.

    ``compute_sha256`` hashes bytes already in memory (the normal ingest path);
    ``stream_sha256`` hashes a file in chunks (the oversized path, which must never load
    it). They are the same digest over the same content — a file grouped one way by one
    and another way by the other would silently break exact-duplicate detection.
    """
    content = b"a" * (3 * 1024 * 1024 + 7)  # spans several chunks, ends mid-chunk
    path = tmp_path / "big.rtf"
    path.write_bytes(content)

    digest, size = stream_sha256(path)
    assert digest == compute_sha256(content)
    assert size == len(content)


def test_stream_sha256_never_loads_the_whole_file(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Constant memory is the entire reason this function exists: it is what lets an
    89 MB document be a ledger row without ever being loaded. ``read_bytes`` explodes,
    so the only way to pass is to stream."""
    path = tmp_path / "huge.rtf"
    path.write_bytes(b"streamed, never slurped")

    def _explode(self: Path) -> bytes:  # pragma: no cover - must never be called
        raise AssertionError(f"{self} was slurped into memory instead of streamed")

    monkeypatch.setattr(Path, "read_bytes", _explode)
    digest, size = stream_sha256(path)
    assert len(digest) == 64
    assert size == len(b"streamed, never slurped")


def _docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for line in paragraphs:
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_txt_tolerant_decode() -> None:
    text = extract_text("Café résumé — Analyst\n".encode(), DocumentFormat.TXT)
    assert "Café résumé" in text
    assert "Analyst" in text


def test_extract_txt_utf16_bom_is_not_read_as_latin1() -> None:
    """latin-1 decodes ANY byte sequence, so a UTF-16 file that is not recognised up
    front does not fail — it silently becomes mojibake. Measured on the archive: 24
    UTF-16LE ``.txt`` exports decoded to ``ÿþP%P%P%…``, losing the entire document."""
    blob = "═════ Job Profiles\nSecretary, grade 7\n".encode("utf-16")
    assert blob[:2] == b"\xff\xfe"  # the BOM that broke the old ladder

    text = extract_text(blob, DocumentFormat.TXT)
    assert "Job Profiles" in text
    assert "Secretary, grade 7" in text
    assert "ÿþ" not in text
    assert "P%" not in text


def test_extract_txt_latin1_fallback() -> None:
    # 0xE9 is 'é' in latin-1 but invalid as a lone UTF-8 byte — the ladder falls
    # through to latin-1 rather than raising.
    text = extract_text(b"Caf\xe9 Analyst", DocumentFormat.TXT)
    assert "Analyst" in text


def test_extract_txt_strips_nul() -> None:
    text = extract_text(b"Research\x00 Analyst\x00", DocumentFormat.TXT)
    assert "\x00" not in text
    assert "Research Analyst" in text


def test_extract_rtf_fixture() -> None:
    text = extract_text((FIXTURES / "sample.rtf").read_bytes(), DocumentFormat.RTF)
    assert "Research Analyst" in text
    assert "\\rtf" not in text  # control words stripped


def test_extract_docx_roundtrip() -> None:
    blob = _docx_bytes(["POSITION SUMMARY", "The Research Analyst compiles data."])
    text = extract_text(blob, DocumentFormat.DOCX)
    assert "POSITION SUMMARY" in text
    assert "Research Analyst compiles data." in text


def test_extract_docx_bad_bytes_raises() -> None:
    with pytest.raises(ExtractionError):
        extract_text(b"not a real docx zip", DocumentFormat.DOCX)


# ── table / content-control recovery (the .docx extraction defect) ──────────
#
# `doc.paragraphs` (python-docx's high-level API) returns only body-level
# `<w:p>` elements — it skips text inside `<w:tbl>` (tables) and `<w:sdt>`
# (Word content controls). Measured on the real archive: 2,596 of 9,947
# `.docx` lose >40% of their text this way, 24 lose everything. These tests
# pin the fix (a document-order walk of `doc.element.body` that descends into
# both) and the blast-radius guard (plain-paragraph docs are unaffected).


def _old_extract_docx_paragraphs_only(blob: bytes) -> str:
    """The OLD (buggy) extractor, kept here as an oracle: paragraphs only, no
    table/SDT descent. Used both to prove the defect (table/SDT text is
    genuinely absent from its output) and as the literal expression the
    byte-identical guard must match on plain documents."""
    doc = Document(BytesIO(blob))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def _docx_bytes_with_table(rows: list[list[str]]) -> bytes:
    doc = Document()
    doc.add_paragraph("before-table")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            table.cell(r, c).text = cell_text
    doc.add_paragraph("after-table")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sdt_paragraph_xml(text: str) -> str:
    """A `<w:sdt>` (content control) wrapping a single paragraph run — the
    shape python-docx has no high-level API for, built directly as OOXML."""
    return (
        f'<w:sdt {nsdecls("w")}><w:sdtContent>'
        f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
        "</w:sdtContent></w:sdt>"
    )


def _docx_bytes_with_sdt(text: str = "sdt-content-text") -> bytes:
    doc = Document()
    doc.add_paragraph("before-sdt")
    sdt = parse_xml(_sdt_paragraph_xml(text))
    doc.element.body.insert(1, sdt)  # between the two paragraphs, before sectPr
    doc.add_paragraph("after-sdt")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_bytes_plain_with_tab_break_and_blank_paragraphs() -> bytes:
    """A plain-paragraph doc — no table, no SDT — with the edge cases the
    guard must reproduce exactly: a tab and a line break inside one run's
    text, an empty paragraph, and a whitespace-only paragraph."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("before-tab")
    run.add_tab()
    run2 = p.add_run("after-tab")
    run2.add_break()
    p.add_run("after-break")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Trailing paragraph.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_bytes_with_nested_table_and_sdt_in_cell() -> bytes:
    """Table-in-cell (nested table) and SDT-in-cell (nested content control) —
    the recursion the fix must apply to arbitrary nesting depth."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    outer_cell = outer.cell(0, 0)
    outer_cell.paragraphs[0].text = "outer-cell-text"
    inner = outer_cell.add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "inner-table-text"

    sdt_table = doc.add_table(rows=1, cols=1)
    sdt_cell = sdt_table.cell(0, 0)
    sdt = parse_xml(_sdt_paragraph_xml("sdt-in-cell-text"))
    sdt_cell._tc.append(
        sdt
    )  # noqa: SLF001 — test-only oxml construction, no public API

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_table_text_is_missed_by_paragraphs_only() -> None:
    """Pins the DEFECT: the old paragraphs-only join genuinely cannot see
    table cell text. If this goes red, the oracle below is not testing what
    it claims to."""
    blob = _docx_bytes_with_table([["cell00", "cell01"], ["cell10", "cell11"]])
    old = _old_extract_docx_paragraphs_only(blob)
    assert "cell00" not in old
    assert "cell11" not in old


def test_extract_docx_table_recovered_in_row_column_order() -> None:
    blob = _docx_bytes_with_table([["cell00", "cell01"], ["cell10", "cell11"]])
    text = extract_text(blob, DocumentFormat.DOCX)
    assert "before-table" in text
    assert "after-table" in text
    for cell_text in ("cell00", "cell01", "cell10", "cell11"):
        assert cell_text in text
    assert (
        text.index("cell00")
        < text.index("cell01")
        < text.index("cell10")
        < text.index("cell11")
    )


def test_extract_docx_content_control_text_is_missed_by_paragraphs_only() -> None:
    """Pins the DEFECT: SDT content is invisible to `doc.paragraphs` — this is
    the `Registered_Nurse.docx` case (~6k chars of real JD text, recovered vs 0
    today, its whole body sitting inside a content control)."""
    blob = _docx_bytes_with_sdt("sdt-content-text")
    old = _old_extract_docx_paragraphs_only(blob)
    assert "sdt-content-text" not in old


def test_extract_docx_content_control_recovered() -> None:
    blob = _docx_bytes_with_sdt("sdt-content-text")
    text = extract_text(blob, DocumentFormat.DOCX)
    assert "before-sdt" in text
    assert "sdt-content-text" in text
    assert "after-sdt" in text


def test_extract_docx_plain_paragraphs_byte_identical_to_old_output() -> None:
    """THE BLAST-RADIUS GUARD. A .docx with no <w:tbl> and no <w:sdt> must
    extract to EXACTLY what the old paragraphs-only expression produced —
    tabs/breaks inside a run, an empty paragraph, and a whitespace-only
    paragraph included. This bounds the fix to only the affected files."""
    blob = _docx_bytes_plain_with_tab_break_and_blank_paragraphs()
    doc = Document(BytesIO(blob))
    old_expression = "\n".join(p.text for p in doc.paragraphs if p.text)

    text = extract_text(blob, DocumentFormat.DOCX)
    assert text == old_expression


def test_extract_docx_nested_table_and_sdt_in_cell_recovered_once_each() -> None:
    """Nesting: table-in-cell and SDT-in-cell must both be recovered, and
    neither double-counted."""
    blob = _docx_bytes_with_nested_table_and_sdt_in_cell()
    text = extract_text(blob, DocumentFormat.DOCX)
    assert text.count("outer-cell-text") == 1
    assert text.count("inner-table-text") == 1
    assert text.count("sdt-in-cell-text") == 1


def test_extract_docx_document_order_preserved() -> None:
    """paragraph, then table, then paragraph -> output preserves that order."""
    blob = _docx_bytes_with_table([["table-cell"]])
    text = extract_text(blob, DocumentFormat.DOCX)
    before, cell, after = (
        text.index("before-table"),
        text.index("table-cell"),
        text.index("after-table"),
    )
    assert before < cell < after


def test_extract_docx_empty_document_extracts_to_empty_string() -> None:
    doc = Document()
    buf = BytesIO()
    doc.save(buf)
    text = extract_text(buf.getvalue(), DocumentFormat.DOCX)
    assert text == ""


def test_extract_docx_excludes_header_prose_and_all_footers() -> None:
    """BODY ONLY, with ONE measured exception — see the identification-block
    tests below.

    The original invariant was *body only, unconditionally*: for this corpus the
    territorial acknowledgement lives in the document body (``document.xml``), not
    in ``footer*.xml``, and the whole validation/HR baseline reads body text — so
    hauling header/footer prose in would move HR numbers for no gain.

    That still holds for **prose**. What it wrongly excluded was the modern SFU
    template's *identification table*, which lives only in ``header*.xml``
    (measured: 4,968 of 9,948 archive ``.docx`` carry ``Position Title:`` in the
    header and nowhere in the body) — the root cause of the 34% paragraph-titles
    defect. So the rule is now: header **prose** is excluded, footers are excluded
    entirely, and a header is admitted *only* when it carries the template's
    identification labels.

    This test pins the exclusion half; ``test_extract_docx_identification_block_*``
    pins the exception.
    """
    doc = Document()
    doc.add_paragraph("BODY_SENTINEL")
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "HEADER_SENTINEL"  # prose, no labels
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = "Position Title: FOOTER_SENTINEL"
    buf = BytesIO()
    doc.save(buf)
    blob = buf.getvalue()

    # The sentinels really are in the document (guards the test itself).
    reloaded = Document(BytesIO(blob))
    assert reloaded.sections[0].header.paragraphs[0].text == "HEADER_SENTINEL"
    assert "FOOTER_SENTINEL" in reloaded.sections[0].footer.paragraphs[0].text

    text = extract_text(blob, DocumentFormat.DOCX)
    assert "BODY_SENTINEL" in text
    assert "HEADER_SENTINEL" not in text
    # A footer is excluded even when it carries an identification label — the
    # corpus keeps only revision dates there, and appending them would land in
    # whichever body section happens to be last (7,171 .docx have footer text).
    assert "FOOTER_SENTINEL" not in text


# ── identification block recovered from the docx header (the 34% defect) ────
#
# Measured over all 14,565 archive files: 4,968 `.docx` carry `Position Title:`
# ONLY in `header*.xml` — zero of them repeat it in the body. Those are exactly
# the documents whose title parsed as a paragraph, because the segmenter found no
# title label and fell back to the first content line (the About-SFU banner or the
# Position Summary prose). The header also carries `Position #:` (4,900),
# `Employee Group:` (4,542), `Department:` (1,380) and `Grade:` (851) — the same
# fields measured as missing/garbage in
# `docs/audit/data-state-and-grade-2026-08-01.md`.


def _docx_bytes_with_header_id_table(
    rows: list[list[str]],
    *,
    body: tuple[str, ...] = ("POSITION SUMMARY", "The Executive Assistant supports."),
    first_page: bool = True,
) -> bytes:
    """A .docx shaped like the modern SFU template: a two-column identification
    table in the (first-page) header, and a body that starts at POSITION SUMMARY."""
    doc = Document()
    for line in body:
        doc.add_paragraph(line)
    section = doc.sections[0]
    if first_page:
        section.different_first_page_header_footer = True
    header = section.first_page_header if first_page else section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6))
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            table.cell(r, c).text = cell_text
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


_ID_ROWS = [
    ["Position Title:", "Executive Assistant to the Vice-President"],
    ["Position #:", "00100537"],
    ["Department:", "VP Advancement and Alumni Engagement"],
    ["Employee Group:", "APSA"],
    ["Grade:", "12"],
]


def test_extract_docx_identification_block_is_invisible_to_a_body_only_walk() -> None:
    """Pins the DEFECT: the identification table is genuinely unreachable from
    ``doc.element.body``. If this goes red the fix below is not testing what it
    claims to."""
    blob = _docx_bytes_with_header_id_table(_ID_ROWS)
    doc = Document(BytesIO(blob))
    body_only = "\n".join(t for t in _iter_docx_block_text(doc.element.body) if t)
    assert "Executive Assistant to the Vice-President" not in body_only
    assert "00100537" not in body_only


def test_extract_docx_identification_block_recovered_as_label_value_lines() -> None:
    """Each header table ROW collapses to ONE ``Label: value`` line — the shape
    the segmenter's identification regexes read. Cell-per-line (what the body walk
    produces) would put the label and its value on separate lines, where those
    regexes cannot see them."""
    text = extract_text(_docx_bytes_with_header_id_table(_ID_ROWS), DocumentFormat.DOCX)
    assert "Position Title: Executive Assistant to the Vice-President" in text
    assert "Position #: 00100537" in text
    assert "Department: VP Advancement and Alumni Engagement" in text
    assert "Employee Group: APSA" in text
    assert "Grade: 12" in text


def test_extract_docx_identification_block_is_marked_and_leads_the_document() -> None:
    """The block is emitted under the canonical IDENTIFICATION heading, first.

    The marker is what scopes the identification-field extraction to this block
    instead of the whole document — the defect behind the 430 garbage ``grade``
    values ('Effective Date: February', 'Assistant'), which were captured by a
    ``Grade:`` regex run over the entire text."""
    text = extract_text(_docx_bytes_with_header_id_table(_ID_ROWS), DocumentFormat.DOCX)
    assert text.startswith(IDENTIFICATION_MARKER + "\n")
    assert text.index("Position Title:") < text.index("POSITION SUMMARY")


def test_extract_docx_identification_block_read_from_a_plain_header_too() -> None:
    """Not every template uses a distinct first-page header."""
    text = extract_text(
        _docx_bytes_with_header_id_table(_ID_ROWS, first_page=False),
        DocumentFormat.DOCX,
    )
    assert "Position Title: Executive Assistant to the Vice-President" in text


def test_extract_docx_running_header_page_furniture_stays_excluded() -> None:
    """A one-label header is page furniture, not an identification block.

    The template's *running* header is ``Position #: <numbers> <page no>``. Where the
    first-page header is absent, admitting that gave 15 documents a position number of
    ``2`` — the page number. A real identification block always carries a title label
    or several fields, so that is what the gate requires."""
    blob = _docx_bytes_with_header_id_table([["Position #:", "2"]], first_page=False)
    doc = Document(BytesIO(blob))
    body_only = "\n".join(t for t in _iter_docx_block_text(doc.element.body) if t)
    assert extract_text(blob, DocumentFormat.DOCX) == body_only


def test_extract_docx_partly_unfilled_identification_block_is_still_admitted() -> None:
    """…but a real block with blank fields IS one. This document (measured) has an
    empty ``Position Title:`` and ``Position #: 00`` — unfilled, not absent, and the
    parse should reflect what it says rather than discard the section."""
    text = extract_text(
        _docx_bytes_with_header_id_table(
            [
                ["Position Title:", ""],
                ["Position #:", "00"],
                ["Employee Group:", "APSA"],
            ]
        ),
        DocumentFormat.DOCX,
    )
    assert text.startswith(IDENTIFICATION_MARKER + "\n")
    assert parse_jd(text).jd.employee_group == "apsa"


def test_extract_docx_plural_position_number_label_captures_the_number() -> None:
    """``Position #s:`` is a real header spelling (16 documents). The label regex read
    the plural ``s`` as the value, so **243 rows** parsed to ``position_number = "s"``
    — the HRIS join key, and visible as-is on the archive browser."""
    text = extract_text(
        _docx_bytes_with_header_id_table(
            [["Position Title:", "Analyst"], ["Position #s:", "00127946, 00127947"]]
        ),
        DocumentFormat.DOCX,
    )
    assert parse_jd(text).jd.position_number == "00127946"


def test_extract_docx_header_without_identification_labels_stays_excluded() -> None:
    """THE BLAST-RADIUS GUARD for the exception: a header that is page furniture
    (a running title, a page number) must not become a bogus IDENTIFICATION
    section. Output must be byte-identical to the body-only walk."""
    blob = _docx_bytes_with_header_id_table(
        [["Executive Assistant", "Page 2 of 4"]], first_page=False
    )
    doc = Document(BytesIO(blob))
    body_only = "\n".join(t for t in _iter_docx_block_text(doc.element.body) if t)
    assert extract_text(blob, DocumentFormat.DOCX) == body_only


def test_extract_docx_no_header_output_unchanged() -> None:
    """A .docx with no header at all is untouched by this change."""
    blob = _docx_bytes(["POSITION SUMMARY", "The Research Analyst compiles data."])
    doc = Document(BytesIO(blob))
    body_only = "\n".join(t for t in _iter_docx_block_text(doc.element.body) if t)
    assert extract_text(blob, DocumentFormat.DOCX) == body_only


def test_extract_docx_identification_row_with_two_label_value_pairs_is_split() -> None:
    """THE REAL ARCHIVE SHAPE. The template packs two pairs onto one row —
    ``Employee Group: | APSA | Grade: | 13`` — which is how **874 of the 876**
    documents that state a grade write it. Joining the whole row would leave
    ``Grade:`` mid-line, invisible to the line-anchored regexes that read it."""
    text = extract_text(
        _docx_bytes_with_header_id_table(
            [
                ["Position Title:", "Portfolio Manager", "Position #:", "00103132"],
                ["Employee Group:", "APSA", "Grade:", "13"],
            ]
        ),
        DocumentFormat.DOCX,
    )
    lines = text.splitlines()
    assert "Position Title: Portfolio Manager" in lines
    assert "Position #: 00103132" in lines
    assert "Employee Group: APSA" in lines
    assert "Grade: 13" in lines

    jd = parse_jd(text).jd
    assert jd.title == "Portfolio Manager"
    assert jd.position_number == "00103132"
    assert jd.classification is not None
    assert (jd.classification.scheme, jd.classification.value) == ("apsa", "13")


def test_extract_docx_identification_unfilled_label_yields_no_grade() -> None:
    """An empty ``Grade:`` cell must stay empty — a grade the document does not
    state is never manufactured (the field is assigned post-authoring for most
    JDFN roles and lives in the HRIS)."""
    text = extract_text(
        _docx_bytes_with_header_id_table(
            [["Position Title:", "Portfolio Manager"], ["Grade:", ""]]
        ),
        DocumentFormat.DOCX,
    )
    assert "Grade:" in text
    assert parse_jd(text).jd.classification is None


def test_extract_docx_identification_block_parses_to_real_fields() -> None:
    """END TO END — the two halves must connect. Extraction shape is only useful
    if ``parse_jd`` reads it, so this asserts the parsed contract, not the text:
    a real title (not a paragraph), and the identification fields the audit
    measured as missing."""
    text = extract_text(_docx_bytes_with_header_id_table(_ID_ROWS), DocumentFormat.DOCX)
    jd = parse_jd(text).jd
    assert jd.title == "Executive Assistant to the Vice-President"
    assert jd.position_number == "00100537"
    assert jd.department == "VP Advancement and Alumni Engagement"
    assert jd.employee_group == "apsa"
    assert jd.classification is not None
    assert jd.classification.value == "12"


def test_extract_legacy_doc_via_antiword() -> None:
    blob = (FIXTURES / "sample_legacy.doc").read_bytes()
    text = extract_text(blob, DocumentFormat.DOC)
    # Stable substrings from the 1982 Systems Analyst II JD (antiword output).
    assert "Systems Analyst" in text
    assert "Computing Centre" in text
    assert "\x00" not in text


def test_unsupported_format_raises() -> None:
    with pytest.raises(UnsupportedFormatError):
        extract_text(b"\x00\x01\x02", DocumentFormat.OTHER)


def test_extract_text_from_path_txt(tmp_path: Path) -> None:
    p = tmp_path / "jd.txt"
    p.write_text("Position Summary: Research Analyst.", encoding="utf-8")
    assert "Research Analyst" in extract_text_from_path(p)


def test_extract_text_from_path_unsupported(tmp_path: Path) -> None:
    p = tmp_path / "scan.tif"
    p.write_bytes(b"\x00\x01")
    with pytest.raises(UnsupportedFormatError):
        extract_text_from_path(p)


# ── DoS hardening ────────────────────────────────────────────────────────────


def test_default_cap_is_50_mib() -> None:
    assert MAX_DOCUMENT_BYTES == 50 * 1024 * 1024


def test_oversized_input_rejected(monkeypatch: pytest.MonkeyPatch) -> None:
    # Patch the cap small so the test needn't allocate 50 MiB.
    monkeypatch.setattr(extract_mod, "MAX_DOCUMENT_BYTES", 8, raising=True)
    with pytest.raises(DocumentTooLargeError):
        extract_text(b"x" * 9, DocumentFormat.TXT)


def test_at_cap_input_allowed(monkeypatch: pytest.MonkeyPatch) -> None:
    # Exactly at the cap is permitted (boundary is strictly greater-than).
    monkeypatch.setattr(extract_mod, "MAX_DOCUMENT_BYTES", 8, raising=True)
    assert extract_text(b"A" * 8, DocumentFormat.TXT).startswith("A")


def test_extract_text_from_path_oversized(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    p = tmp_path / "huge.txt"
    p.write_bytes(b"small on disk")
    # Report an over-cap size without writing a 50 MiB file to disk.
    monkeypatch.setattr(extract_mod, "MAX_DOCUMENT_BYTES", 4, raising=True)
    with pytest.raises(DocumentTooLargeError):
        extract_text_from_path(p)


def test_antiword_timeout_raises_extraction_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def _raise_timeout(*args: object, **kwargs: object) -> None:
        raise subprocess.TimeoutExpired(cmd="antiword", timeout=1)

    monkeypatch.setattr(extract_mod.subprocess, "run", _raise_timeout)
    with pytest.raises(ExtractionError, match="timed out"):
        extract_text(b"fake .doc bytes", DocumentFormat.DOC)
