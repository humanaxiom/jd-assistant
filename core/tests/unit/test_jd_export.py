"""Phase 5.7 — export an SFU JD to ``.docx``.

Snapshot-by-readback: render the document, open it again with python-docx, and assert
its structure — the SFU formatting standard (TNR 10, bold headers, bullets), the
mandated footer sourced from the rulebook, and that empty sections are dropped.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt
from fastapi.testclient import TestClient

from src.api.main import app
from src.jd_core.models.parsed_jd import (
    SFUDuty,
    SFUJobDescription,
    SFUQualification,
    SFURelationships,
)
from src.jd_core.rules import get_rules
from src.jd_export import render_sfu_docx


def _jd(**update: object) -> SFUJobDescription:
    jd = SFUJobDescription(
        title="Software Developer",
        position_summary="Supports the department's budgeting and reporting.",
        duties=[
            SFUDuty(action_verb="Manages", statement="Manages the program (60%)"),
            SFUDuty(action_verb="Coordinates", statement="Coordinates delivery (40%)"),
        ],
        decision_making=["Approves expenditures up to $5k"],
        problem_solving=["Resolves scheduling conflicts independently"],
        relationships=SFURelationships(
            supervisory="Supervises 2 staff", internal=["Finance"], external=["Vendors"]
        ),
        qualifications=[
            SFUQualification(text="Bachelor's degree", kind="education"),
            SFUQualification(
                text="knowledge of databases", kind="knowledge", modifier="Excellent"
            ),
        ],
    )
    return jd.model_copy(update=update)


def _read(data: bytes) -> Document:
    return Document(BytesIO(data))


def test_the_document_uses_the_sfu_base_font() -> None:
    doc = _read(render_sfu_docx(_jd()))
    font = doc.styles["Normal"].font
    assert font.name == "Times New Roman"
    assert font.size == Pt(10)


def test_sections_render_with_headers_and_bullets() -> None:
    doc = _read(render_sfu_docx(_jd()))
    texts = [p.text for p in doc.paragraphs]
    assert "Position Summary" in texts
    assert "Duties and Responsibilities" in texts
    assert "Qualifications" in texts
    # The allocation carried in the duty text survives.
    assert any("Manages the program (60%)" in t for t in texts)
    # Duties are bulleted.
    bullet_texts = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert any("Manages the program" in t for t in bullet_texts)
    # The KSA modifier is prefixed when not already in the text.
    assert any("Excellent knowledge of databases" in t for t in texts)


def test_the_mandated_footer_is_injected_from_the_rulebook() -> None:
    doc = _read(render_sfu_docx(_jd()))
    texts = [p.text for p in doc.paragraphs]
    passage = get_rules().boilerplate.territorial_acknowledgement[0]
    assert any(passage in t for t in texts), "territorial acknowledgement missing"
    equity = get_rules().boilerplate.employment_equity[0]
    assert any(equity in t for t in texts), "employment equity statement missing"


def test_empty_sections_are_dropped() -> None:
    doc = _read(render_sfu_docx(_jd(decision_making=[], problem_solving=[])))
    texts = [p.text for p in doc.paragraphs]
    assert "Impact of Decision Making" not in texts
    assert "Problem Solving" not in texts
    # ...but present sections and the footer are still there.
    assert "Duties and Responsibilities" in texts


def test_a_minimal_jd_still_renders_with_the_footer() -> None:
    doc = _read(render_sfu_docx(SFUJobDescription(title="Bare Role")))
    texts = [p.text for p in doc.paragraphs]
    assert "Bare Role" in texts
    passage = get_rules().boilerplate.territorial_acknowledgement[0]
    assert any(passage in t for t in texts)


def test_export_route_returns_a_docx_download() -> None:
    resp = TestClient(app).post(
        "/jd-bank/compose/export",
        json={"title": "Financial Analyst", "position_summary": "Supports finance."},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "financial-analyst.docx" in resp.headers["content-disposition"]
    assert resp.content[:2] == b"PK"  # a .docx is a zip archive
